from docx import Document
import socket

START_PORT = 1
END_PORT = 500

document = Document()
document.add_heading("Open ports", 0)


def basic_scan(target):

    open_ports = []

    for port in range(START_PORT, END_PORT + 1):
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(0.2)
        result = sock.connect_ex((target, port))
        if result == 0:
            open_ports.append(port)
            print(f"\033[1;32m Port {port} is open\33[0m")
        print(port)



        sock.close()

    print("Scan finished.")

    for port in open_ports:
        document.add_paragraph(f"Port {port} is open")

    document.save("open_ports.docx")

basic_scan("192.168.1.1")

